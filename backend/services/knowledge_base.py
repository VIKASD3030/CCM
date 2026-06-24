"""
Phase 1 — Knowledge Base Service.
Handles document upload, text extraction, chunking, embedding, and vector search.
PostgreSQL 16 with in-Python cosine similarity for vector search.
"""
import uuid
import logging

from backend.config import get_settings
from backend.models.document import KnowledgeDocument, DocumentChunk
from backend.models.audit import AuditEntry
from backend.services.webhook_service import dispatch_event
from backend.services.storage import storage_backend
from backend.services import gemini as gemini_client

from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)
settings = get_settings()


async def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from a file. Supports PDF, DOCX, and Images."""
    file_type = file_type.lower()
    if file_type == "pdf":
        try:
            import pymupdf4llm
            return pymupdf4llm.to_markdown(file_path)
        except Exception as e:
            logger.warning(f"pymupdf4llm failed, falling back to pymupdf: {e}")
            import pymupdf
            doc = pymupdf.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return text
    elif file_type == "docx":
        import docx
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif file_type in ("jpg", "jpeg", "png"):
        import base64
        mime_type = "image/jpeg" if file_type in ("jpg", "jpeg") else "image/png"
        with open(file_path, "rb") as image_file:
            image_bytes = image_file.read()

        response = await gemini_client.call_gemini_vision_async(
            text="Extract all text from this image accurately. Return only the extracted text.",
            image_data=image_bytes,
            mime_type=mime_type,
            temperature=0.1,
            max_output_tokens=4096,
        )
        return response
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Supported types: pdf, docx, jpg, jpeg, png")


async def extract_text_from_bytes(content: bytes, file_type: str, filename: str) -> str:
    """Extract text from file bytes without saving to disk first."""
    import tempfile
    import os

    suffix = f".{file_type}"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        return await extract_text_from_file(tmp_path, file_type)
    finally:
        os.unlink(tmp_path)


def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> list[str]:
    """
    Split text into overlapping chunks by word count.
    Returns list of chunk strings.
    """
    chunk_size = chunk_size or settings.chunk_size
    overlap = overlap or settings.chunk_overlap

    words = text.split()
    if not words:
        return []

    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start = end - overlap
        if start >= len(words):
            break

    return chunks


async def generate_embeddings(texts: list[str]) -> list[list[float]]:
    """Generate embeddings in batches with error handling."""
    embeddings = []
    batch_size = 100
    for i in range(0, len(texts), batch_size):
        batch = texts[i: i + batch_size]
        try:
            batch_embeddings = await gemini_client.generate_embedding_async(batch)
            for emb in batch_embeddings:
                embeddings.append(emb)
        except Exception as e:
            logger.error(f"Embedding generation failed for batch {i}: {e}")
            embeddings.extend([None] * len(batch))

    return embeddings


async def init_document(
    db: AsyncSession,
    filename: str,
    content: bytes,
    file_type: str,
    category: str = "general",
    user_id: uuid.UUID = None,
) -> KnowledgeDocument:
    """
    Create a document record and save the file to storage.
    Heavy work (chunking, embedding) happens in the background ARQ job.
    """
    # Save file to storage
    import uuid as uuid_gen
    storage_key = f"kb/{uuid_gen.uuid4()}/{filename}"
    await storage_backend.upload(content, storage_key, f"application/{file_type}")

    doc = KnowledgeDocument(
        filename=filename,
        file_type=file_type,
        category=category,
        status="processing",
        storage_path=storage_key,
        uploaded_by=user_id,
    )
    db.add(doc)
    await db.flush()

    logger.info("kb.document.initialized", document_id=str(doc.id), filename=filename)
    return doc


async def ingest_document(
    db: AsyncSession,
    filename: str,
    content: bytes,
    file_type: str,
    category: str = "general",
    user_id: uuid.UUID = None,
) -> KnowledgeDocument:
    """
    Full ingestion pipeline: extract text -> chunk -> embed -> store.
    Dispatches webhook on completion.
    """
    doc = KnowledgeDocument(
        filename=filename,
        file_type=file_type,
        category=category,
        status="processing",
        uploaded_by=user_id,
    )
    db.add(doc)
    await db.flush()

    try:
        raw_text = await extract_text_from_bytes(content, file_type, filename)
        doc.raw_text = raw_text

        chunks = chunk_text(raw_text)
        if not chunks:
            doc.status = "failed"
            doc.chunk_count = 0
            await db.flush()
            return doc

        embeddings = await generate_embeddings(chunks)

        for idx, (chunk_text_str, embedding) in enumerate(zip(chunks, embeddings)):
            chunk = DocumentChunk(
                document_id=str(doc.id),
                chunk_index=idx,
                chunk_text=chunk_text_str,
                embedding=embedding,
                metadata_={
                    "filename": filename,
                    "category": category,
                    "chunk_index": idx,
                    "total_chunks": len(chunks),
                },
            )
            db.add(chunk)

        doc.chunk_count = len(chunks)
        doc.status = "indexed"

        audit = AuditEntry(
            entity_type="document",
            entity_id=str(doc.id),
            action="uploaded_and_indexed",
            user_id=user_id,
            details={
                "filename": filename,
                "category": category,
                "chunks": len(chunks),
            },
        )
        db.add(audit)

        await db.flush()
        logger.info(f"Ingested document '{filename}': {len(chunks)} chunks indexed by user {user_id}")

        # Dispatch webhook (fire-and-forget)
        await dispatch_event("kb.document_added", {
            "document_id": str(doc.id),
            "filename": filename,
            "category": category,
            "chunks": len(chunks),
            "uploaded_by": str(user_id) if user_id else None,
        })

    except Exception as e:
        doc.status = "failed"
        logger.error(f"Failed to ingest document '{filename}': {e}")
        raise

    return doc


def cosine_similarity(a: list[float], b: list[float]) -> float:
    """Compute cosine similarity between two vectors."""
    import math
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x * x for x in a))
    norm_b = math.sqrt(sum(x * x for x in b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


async def search_similar_chunks(
    db: AsyncSession,
    query_text: str,
    top_k: int = 5,
    category: str = None,
    user_id: uuid.UUID = None,
    user_role: str = None,
) -> list[dict]:
    """
    Semantic search using in-Python cosine similarity over stored embeddings.
    """
    query_embeddings = await generate_embeddings([query_text])
    if not query_embeddings or query_embeddings[0] is None:
        return []

    query_vec = query_embeddings[0]

    stmt = select(DocumentChunk).join(
        KnowledgeDocument, DocumentChunk.document_id == KnowledgeDocument.id
    ).where(KnowledgeDocument.status == "indexed")

    if user_role not in ["admin", "reviewer"] and user_id is not None:
        stmt = stmt.where(KnowledgeDocument.uploaded_by == user_id)

    if category:
        stmt = stmt.where(KnowledgeDocument.category == category)

    result = await db.execute(stmt)
    chunks = result.scalars().all()

    scored = []
    for chunk in chunks:
        if chunk.embedding is None:
            continue
        sim = cosine_similarity(query_vec, chunk.embedding)
        scored.append((sim, chunk))

    scored.sort(key=lambda x: x[0], reverse=True)
    scored = scored[:top_k]

    return [
        {
            "chunk_id": str(chunk.id),
            "document_id": str(chunk.document_id),
            "chunk_text": chunk.chunk_text,
            "chunk_index": chunk.chunk_index,
            "similarity": round(sim, 4),
            "metadata": chunk.metadata_ if chunk.metadata_ else {},
        }
        for sim, chunk in scored
    ]


async def get_all_documents(
    db: AsyncSession,
    user_id: uuid.UUID = None,
    user_role: str = None,
    skip: int = 0,
    limit: int = 50,
) -> tuple[list[dict], int]:
    """List all knowledge base documents with pagination and row-level security."""
    count_stmt = select(func.count()).select_from(KnowledgeDocument)
    base_stmt = select(KnowledgeDocument)

    if user_role not in ["admin", "reviewer"] and user_id is not None:
        count_stmt = count_stmt.where(KnowledgeDocument.uploaded_by == user_id)
        base_stmt = base_stmt.where(KnowledgeDocument.uploaded_by == user_id)

    total_result = await db.execute(count_stmt)
    total = total_result.scalar() or 0

    stmt = base_stmt.order_by(KnowledgeDocument.uploaded_at.desc()).offset(skip).limit(limit)
    result = await db.execute(stmt)
    documents = result.scalars().all()
    return [doc.to_dict() for doc in documents], total


async def get_document_stats(
    db: AsyncSession,
    user_id: uuid.UUID = None,
    user_role: str = None,
) -> dict:
    """Get knowledge base statistics with row-level security."""
    doc_count_stmt = select(func.count()).select_from(KnowledgeDocument)
    chunk_count_stmt = select(func.count()).select_from(DocumentChunk)
    indexed_count_stmt = (
        select(func.count())
        .select_from(KnowledgeDocument)
        .where(KnowledgeDocument.status == "indexed")
    )
    category_stmt = (
        select(KnowledgeDocument.category, func.count())
        .group_by(KnowledgeDocument.category)
    )

    if user_role not in ["admin", "reviewer"] and user_id is not None:
        doc_count_stmt = doc_count_stmt.where(KnowledgeDocument.uploaded_by == user_id)
        chunk_count_stmt = chunk_count_stmt.join(KnowledgeDocument).where(KnowledgeDocument.uploaded_by == user_id)
        indexed_count_stmt = indexed_count_stmt.where(KnowledgeDocument.uploaded_by == user_id)
        category_stmt = category_stmt.where(KnowledgeDocument.uploaded_by == user_id)

    doc_count = await db.execute(doc_count_stmt)
    chunk_count = await db.execute(chunk_count_stmt)
    indexed_count = await db.execute(indexed_count_stmt)
    category_result = await db.execute(category_stmt)

    return {
        "total_documents": doc_count.scalar() or 0,
        "total_chunks": chunk_count.scalar() or 0,
        "indexed_documents": indexed_count.scalar() or 0,
        "categories": {row[0]: row[1] for row in category_result.all()},
    }


async def delete_document(
    db: AsyncSession,
    document_id,
    user_id: uuid.UUID = None,
    user_role: str = None,
) -> bool:
    """Delete a document and all its chunks with row-level security."""
    doc = await db.get(KnowledgeDocument, str(document_id))
    if not doc:
        return False

    if user_role not in ["admin"] and user_id is not None:
        if doc.uploaded_by != user_id:
            return False

    await db.delete(doc)

    audit = AuditEntry(
        entity_type="document",
        entity_id=str(document_id),
        action="deleted",
        user_id=user_id,
        details={"filename": doc.filename},
    )
    db.add(audit)
    await db.flush()
    return True
